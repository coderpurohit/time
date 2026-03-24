"""
Inspect the last uploaded DOCX - write to file for full output.
"""
import docx

doc = docx.Document("last_uploaded_debug.docx")

with open("docx_structure.txt", "w", encoding="utf-8") as out:
    out.write(f"Number of tables: {len(doc.tables)}\n")
    
    for ti, table in enumerate(doc.tables):
        out.write(f"\n{'='*100}\n")
        out.write(f"TABLE {ti}: Rows={len(table.rows)}, Cols={len(table.columns)}\n")
        out.write(f"{'='*100}\n")
        
        for ri, row in enumerate(table.rows):
            cells = []
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip().replace('\n', '\\n')
                cells.append(f"[{ci}]=\"{text}\"")
            out.write(f"  R{ri:02d}: {' | '.join(cells)}\n")

print("Written to docx_structure.txt")
