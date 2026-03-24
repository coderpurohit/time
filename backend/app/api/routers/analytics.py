from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from typing import List, Dict, Any
import io
import docx
import re
from ...infrastructure.database import get_db
from ...infrastructure import models
from .. import schemas
router = APIRouter()

@router.get("/dashboard-stats")
def get_dashboard_statistics(db: Session = Depends(get_db)):
    """
    Get dashboard statistics for the main dashboard
    """
    try:
        # Count all entities
        teachers_count = db.query(models.Teacher).count()
        classes_count = db.query(models.ClassGroup).count()
        rooms_count = db.query(models.Room).count()
        subjects_count = db.query(models.Subject).count()
        timeslots_count = db.query(models.TimeSlot).count()
        
        # Get latest timetable for utilization
        latest_version = db.query(models.TimetableVersion).order_by(
            models.TimetableVersion.id.desc()
        ).first()
        
        utilization = 0
        if latest_version and latest_version.entries:
            total_entries = len(latest_version.entries)
            # Calculate utilization as percentage of possible slots used
            max_possible = timeslots_count * classes_count if timeslots_count and classes_count else 1
            utilization = min(100, round((total_entries / max_possible) * 100)) if max_possible > 0 else 0
        
        return {
            "teachers": teachers_count,
            "classes": classes_count,
            "rooms": rooms_count,
            "subjects": subjects_count,
            "timeSlots": timeslots_count,
            "utilization": utilization,
            "status": "success"
        }
        
    except Exception as e:
        # Return fallback values if there's an error
        return {
            "teachers": 15,
            "classes": 9,
            "rooms": 10,
            "subjects": 12,
            "timeSlots": 35,
            "utilization": 85,
            "status": "fallback",
            "error": str(e)
        }

@router.get("/load-factor")
def get_load_factor_analysis(db: Session = Depends(get_db)):
    """
    Get comprehensive load factor analysis with validation and conflict detection.
    Works even before a timetable is generated — returns teacher list from DB.
    """
    # Get latest timetable
    latest_version = db.query(models.TimetableVersion).order_by(
        models.TimetableVersion.id.desc()
    ).first()
    
    # Check for standalone override (version_id=NULL) if no timetable exists
    if not latest_version:
        standalone_override = db.query(models.WorkloadReportOverride).filter(
            models.WorkloadReportOverride.version_id == None
        ).first()
        
        teachers = db.query(models.Teacher).all()
        classes = db.query(models.ClassGroup).all()
        subjects = db.query(models.Subject).all()
        
        if standalone_override and standalone_override.report_data:
            teacher_load = standalone_override.report_data
        else:
            # Build basic teacher data from DB (no timetable entries)
            teacher_load = []
            for teacher in teachers:
                teacher_load.append({
                    "id": teacher.id,
                    "name": teacher.name,
                    "email": teacher.email,
                    "theory_hours": 0,
                    "practical_hours": 0,
                    "lab_hours": 0,
                    "project_hours": 0,
                    "total_periods": 0,
                    "max_hours": teacher.max_hours_per_week or 20,
                    "load_percentage": 0,
                    "status": "No Timetable",
                    "daily_load": {},
                    "subject_distribution": {},
                    "courses_detail": [],
                    "avg_per_day": 0
                })
        
        return {
            "summary": {
                "total_teachers": len(teachers),
                "total_classes": len(classes),
                "total_periods": 0,
                "avg_teacher_load": 0,
                "classroom_utilization": 0,
                "lab_utilization": 0,
                "timetable_id": 0,
                "timetable_name": "Standalone (No Timetable)"
            },
            "teacher_load": teacher_load,
            "class_load": [],
            "conflicts": {
                "faculty_timing_clashes": [],
                "classroom_conflicts": [],
                "lab_conflicts": [],
                "missing_entries": [],
                "overloaded_faculty": [],
                "underloaded_faculty": []
            },
            "validation": {
                "total_conflicts": 0,
                "critical_issues": 0,
                "warnings": 0
            },
            "subjects": [{"id": s.id, "name": s.name, "code": s.code, "is_lab": s.is_lab} for s in subjects]
        }
    
    entries = latest_version.entries or []
    
    if not entries:
        # Return empty structure instead of 404
        teachers = db.query(models.Teacher).all()
        classes = db.query(models.ClassGroup).all()
        subjects = db.query(models.Subject).all()
        return {
            "summary": {
                "total_teachers": len(teachers),
                "total_classes": len(classes),
                "total_periods": 0,
                "avg_teacher_load": 0,
                "classroom_utilization": 0,
                "lab_utilization": 0,
                "timetable_id": latest_version.id,
                "timetable_name": latest_version.name
            },
            "teacher_load": [],
            "class_load": [],
            "conflicts": {
                "faculty_timing_clashes": [],
                "classroom_conflicts": [],
                "lab_conflicts": [],
                "missing_entries": [],
                "overloaded_faculty": [],
                "underloaded_faculty": []
            },
            "validation": {
                "total_conflicts": 0,
                "critical_issues": 0,
                "warnings": 0
            },
            "subjects": [{"id": s.id, "name": s.name, "code": s.code, "is_lab": s.is_lab} for s in subjects]
        }
    
    # Get all data
    teachers = db.query(models.Teacher).all()
    classes = db.query(models.ClassGroup).all()
    subjects = db.query(models.Subject).all()
    slots = db.query(models.TimeSlot).all()
    rooms = db.query(models.Room).all()
    
    # Initialize conflict tracking
    conflicts = {
        "faculty_timing_clashes": [],
        "classroom_conflicts": [],
        "lab_conflicts": [],
        "missing_entries": [],
        "overloaded_faculty": [],
        "underloaded_faculty": []
    }
    
    # Calculate enhanced teacher load with validation
    teacher_load = []
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday']
    
    for teacher in teachers:
        teacher_entries = [e for e in entries if e.teacher_id == teacher.id]
        
        # Categorize hours by type
        theory_hours = 0
        practical_hours = 0
        lab_hours = 0
        project_hours = 0
        
        # Track timing conflicts for this teacher
        teacher_schedule = {}
        
        # Course-level tracking for detailed load distribution
        course_groups = {}
        
        class_group_entries = {}
        for entry in teacher_entries:
            if not entry.subject: continue
            cls_id = entry.class_group_id if entry.class_group else "all"
            if cls_id not in class_group_entries:
                class_group_entries[cls_id] = []
            class_group_entries[cls_id].append(entry)
            
            # Record timing conflicts
            if entry.time_slot:
                slot_key = f"{entry.time_slot.day}_{entry.time_slot.start_time}"
                if slot_key in teacher_schedule:
                    conflicts["faculty_timing_clashes"].append({
                        "teacher": teacher.name,
                        "day": entry.time_slot.day,
                        "time": entry.time_slot.start_time,
                        "subjects": [teacher_schedule[slot_key], entry.subject.name],
                        "classes": [entry.class_group.name if entry.class_group else "Unknown"]
                    })
                else:
                    teacher_schedule[slot_key] = entry.subject.name
                    
        for cls_id, cls_entries in class_group_entries.items():
            cls_name = cls_entries[0].class_group.name if cls_entries[0].class_group else "All"
            
            theory_list = []
            prac_list = []
            proj_list = []
            
            for entry in cls_entries:
                subject_name = entry.subject.name.lower()
                lab_keywords = ['sl1', 'sl2', 'sl3', 'sl4', 'cl1', 'cl2', 'lab', 'laboratory']
                practical_keywords = ['practical', 'workshop', 'tutorial', 'seminar', 'training', 'hands-on', 'exercise', 'mini project']
                project_keywords = ['project', 'major project', 'capstone', 'dissertation', 'thesis']
                
                is_lab_cat = entry.subject.is_lab or any(kw in subject_name for kw in lab_keywords)
                is_prac_cat = any(kw in subject_name for kw in practical_keywords)
                is_proj_cat = any(kw in subject_name for kw in project_keywords)
                
                duration = 1  # Standard single time slot is 1 hour
                
                # Categorize entry and compute real period count
                if is_proj_cat and "mini" not in subject_name:
                    proj_list.append(entry)
                    project_hours += duration
                elif is_lab_cat or is_prac_cat:
                    prac_list.append(entry)
                    if is_lab_cat: lab_hours += duration
                    else: practical_hours += duration
                else:
                    theory_list.append(entry)
                    theory_hours += duration
            
            # Determine Theory courses for this division
            theory_subs = list({e.subject.id: e.subject for e in theory_list}.values())
            
            primary_subs = theory_subs
            if not primary_subs:
                primary_subs = list({e.subject.id: e.subject for e in prac_list}.values())
                if not primary_subs:
                    primary_subs = list({e.subject.id: e.subject for e in proj_list}.values())
                    
            if primary_subs:
                # Merge all practicals into the first primary subject course
                primary_sub = primary_subs[0]
                
                # Calculate practical hours syntax (Batches = total duration slots / batch size)
                lab_counts = {}
                for e in prac_list:
                    sid = e.subject.id
                    if sid not in lab_counts:
                        duration = e.subject.duration_slots or 1 
                        name = e.subject.name.replace(" Lab", "").replace(" Laboratory", "")
                        lab_counts[sid] = {"name": name, "duration": duration, "count": 0}
                    lab_counts[sid]["count"] += 1
                    
                prac_strs = []
                t_count = len([e for e in theory_list if e.subject.id == primary_sub.id])
                p_count = len(prac_list)
                proj_count = len(proj_list)
                total_row_load = t_count + p_count + proj_count
                
                for lab in lab_counts.values():
                    dur = lab["duration"]
                    batches = lab["count"] // dur if dur > 0 else lab["count"]
                    if batches == 0: batches = 1
                    prac_strs.append(f"{lab['name'] + ' ' if lab['name'] != primary_sub.name else ''}{dur} * {batches}")
                    
                prac_str = "<br>".join(prac_strs) if prac_strs else "-"
                
                course_groups[f"{primary_sub.id}_{cls_id}"] = {
                    "course_name": primary_sub.name,
                    "class_division": cls_name,
                    "theory_hours": t_count if t_count > 0 else "-",
                    "practical_hours": prac_str,
                    "project_hours": "2 * 1",
                    "total_load": total_row_load
                }
                
                # For any additional courses, just populate theory to avoid double counting practicals
                for extra_sub in primary_subs[1:]:
                    e_t_count = len([e for e in theory_list if e.subject.id == extra_sub.id])
                    course_groups[f"{extra_sub.id}_{cls_id}"] = {
                        "course_name": extra_sub.name,
                        "class_division": cls_name,
                        "theory_hours": e_t_count if e_t_count > 0 else "-",
                        "practical_hours": "-",
                        "project_hours": "2 * 1",
                        "total_load": e_t_count
                    }

        courses_detail = list(course_groups.values())
        
        # Merge custom workload aspects
        for cw in teacher.custom_workloads:
            courses_detail.append({
                "id": cw.id,
                "course_name": cw.course_name,
                "class_division": cw.class_division,
                "theory_hours": cw.theory_hours,
                "practical_hours": cw.practical_hours,
                "project_hours": cw.project_hours,
                "total_load": cw.total_load,
                "is_custom": True
            })
            theory_hours += cw.total_load
            
        courses_detail.sort(key=lambda x: (x["course_name"], x["class_division"]))
        
        total_periods = theory_hours + practical_hours + lab_hours + project_hours
        max_hours = teacher.max_hours_per_week or 20
        load_percentage = round((total_periods / max_hours) * 100) if max_hours > 0 else 0
        
        # Determine status
        status = "OK"
        if load_percentage > 100:
            status = "Overloaded"
            conflicts["overloaded_faculty"].append({
                "teacher": teacher.name,
                "total_load": total_periods,
                "max_hours": max_hours,
                "excess": total_periods - max_hours
            })
        elif load_percentage < 50:
            status = "Underloaded"
            conflicts["underloaded_faculty"].append({
                "teacher": teacher.name,
                "total_load": total_periods,
                "max_hours": max_hours,
                "utilization": load_percentage
            })
        
        # Daily distribution
        daily_load = {}
        for day in days:
            day_entries = [e for e in teacher_entries if e.time_slot and e.time_slot.day == day]
            daily_load[day] = len(day_entries)
        
        # Subject distribution
        subject_distribution = {}
        for entry in teacher_entries:
            subject_id = entry.subject_id
            subject_name = entry.subject.name if entry.subject else "Unknown"
            if subject_name not in subject_distribution:
                subject_distribution[subject_name] = 0
            subject_distribution[subject_name] += 1
        
        teacher_load.append({
            "id": teacher.id,
            "name": teacher.name,
            "email": teacher.email,
            "theory_hours": theory_hours,
            "practical_hours": practical_hours,
            "lab_hours": lab_hours,
            "project_hours": project_hours,
            "total_periods": total_periods,
            "max_hours": max_hours,
            "load_percentage": load_percentage,
            "status": status,
            "daily_load": daily_load,
            "subject_distribution": subject_distribution,
            "courses_detail": courses_detail,
            "avg_per_day": round(total_periods / 5, 1)
        })
    
    # Check classroom conflicts
    room_schedule = {}
    for entry in entries:
        if entry.room and entry.time_slot:
            slot_key = f"{entry.time_slot.day}_{entry.time_slot.start_time}_{entry.room_id}"
            
            if slot_key in room_schedule:
                conflicts["classroom_conflicts"].append({
                    "room": entry.room.name,
                    "day": entry.time_slot.day,
                    "time": entry.time_slot.start_time,
                    "classes": [room_schedule[slot_key], entry.class_group.name if entry.class_group else "Unknown"]
                })
            else:
                room_schedule[slot_key] = entry.class_group.name if entry.class_group else "Unknown"
    
    # Check lab conflicts (rooms with type 'Lab')
    lab_schedule = {}
    for entry in entries:
        if entry.room and entry.room.type == 'Lab' and entry.time_slot:
            slot_key = f"{entry.time_slot.day}_{entry.time_slot.start_time}_{entry.room_id}"
            
            if slot_key in lab_schedule:
                conflicts["lab_conflicts"].append({
                    "lab": entry.room.name,
                    "day": entry.time_slot.day,
                    "time": entry.time_slot.start_time,
                    "batches": [lab_schedule[slot_key], entry.class_group.name if entry.class_group else "Unknown"]
                })
            else:
                lab_schedule[slot_key] = entry.class_group.name if entry.class_group else "Unknown"
    
    # Calculate class load with validation
    class_load = []
    max_periods_per_class = 30
    
    for class_group in classes:
        class_entries = [e for e in entries if e.class_group_id == class_group.id]
        total_periods = len(class_entries)
        load_percentage = round((total_periods / max_periods_per_class) * 100) if max_periods_per_class > 0 else 0
        
        # Daily distribution
        daily_load = {}
        for day in days:
            day_entries = [e for e in class_entries if e.time_slot and e.time_slot.day == day]
            daily_load[day] = len(day_entries)
        
        # Subject distribution with hours validation
        subject_distribution = {}
        theory_total = 0
        practical_total = 0
        lab_total = 0
        
        for entry in class_entries:
            subject_name = entry.subject.name if entry.subject else "Unknown"
            if subject_name not in subject_distribution:
                subject_distribution[subject_name] = 0
            subject_distribution[subject_name] += 1
            
            # Enhanced count by type using lab names and keywords
            if entry.subject:
                subject_name = entry.subject.name.lower()
                
                # Lab classification based on your lab timetable
                lab_keywords = [
                    'sl1 lab', 'sl2 lab', 'sl3 lab', 'sl4 lab',  # Software Labs
                    'ai lab', 'ds lab', 'pll lab', 'diy lab',   # AI/DS/Programming Labs
                    'lab', 'laboratory'                          # General lab keywords
                ]
                
                # Practical classification
                practical_keywords = [
                    'practical', 'workshop', 'tutorial', 'seminar',
                    'training', 'hands-on', 'exercise'
                ]
                
                if entry.subject.is_lab or any(keyword in subject_name for keyword in lab_keywords):
                    lab_total += 1
                elif any(keyword in subject_name for keyword in practical_keywords):
                    practical_total += 1
                else:
                    theory_total += 1
        
        class_load.append({
            "id": class_group.id,
            "name": class_group.name,
            "student_count": class_group.student_count,
            "total_periods": total_periods,
            "theory_periods": theory_total,
            "practical_periods": practical_total,
            "lab_periods": lab_total,
            "load_percentage": load_percentage,
            "daily_load": daily_load,
            "subject_distribution": subject_distribution
        })
    
    # Calculate resource utilization
    total_classroom_slots = len([r for r in rooms if r.type != 'Lab']) * len(slots)
    used_classroom_slots = len([e for e in entries if e.room and e.room.type != 'Lab'])
    classroom_utilization = round((used_classroom_slots / total_classroom_slots) * 100) if total_classroom_slots > 0 else 0
    
    total_lab_slots = len([r for r in rooms if r.type == 'Lab']) * len(slots)
    used_lab_slots = len([e for e in entries if e.room and e.room.type == 'Lab'])
    lab_utilization = round((used_lab_slots / total_lab_slots) * 100) if total_lab_slots > 0 else 0
    
    # Check for Workload JSON Override
    override = db.query(models.WorkloadReportOverride).filter(
        models.WorkloadReportOverride.version_id == latest_version.id
    ).first()
    
    if override and override.report_data:
        teacher_load = override.report_data

    # Calculate summary stats
    avg_teacher_load = round(sum(t["total_periods"] for t in teacher_load) / len(teacher_load)) if teacher_load else 0
    
    return {
        "summary": {
            "total_teachers": len(teachers),
            "total_classes": len(classes),
            "total_periods": len(entries),
            "avg_teacher_load": avg_teacher_load,
            "classroom_utilization": classroom_utilization,
            "lab_utilization": lab_utilization,
            "timetable_id": latest_version.id,
            "timetable_name": latest_version.name
        },
        "teacher_load": teacher_load,
        "class_load": class_load,
        "conflicts": conflicts,
        "validation": {
            "total_conflicts": sum(len(v) for v in conflicts.values()),
            "critical_issues": len(conflicts["faculty_timing_clashes"]) + len(conflicts["classroom_conflicts"]),
            "warnings": len(conflicts["overloaded_faculty"]) + len(conflicts["underloaded_faculty"])
        },
        "subjects": [{"id": s.id, "name": s.name, "code": s.code, "is_lab": s.is_lab} for s in subjects]
    }

@router.get("/lab-schedule-analysis")
def get_lab_schedule_analysis(db: Session = Depends(get_db)):
    """
    Get detailed lab schedule analysis based on D.Y. Patil College lab format
    """
    # Get latest timetable
    latest_version = db.query(models.TimetableVersion).order_by(
        models.TimetableVersion.id.desc()
    ).first()
    
    if not latest_version:
        raise HTTPException(status_code=404, detail="No timetable found. Please generate a timetable first.")
    
    entries = latest_version.entries
    subjects = db.query(models.Subject).all()
    rooms = db.query(models.Room).all()
    
    # Define lab categories based on your timetable
    lab_categories = {
        "Software Labs": ["SL1 Lab", "SL2 Lab", "SL3 Lab", "SL4 Lab"],
        "AI/DS Labs": ["AI Lab", "DS Lab"],
        "Programming Labs": ["PLL Lab", "DIY Lab"],
        "General Labs": ["Lab"]  # Catch-all for other labs
    }
    
    # Time slots from your timetable
    time_slots = [
        "8.15-9.15", "9.15-10.15", "10.30-11.30", "11.30-12.30",
        "12.30-1.30", "1.30-2.30", "2.30-3.30", "3.45-4.45", "4.45-5.45"
    ]
    
    # Analyze lab utilization
    lab_utilization = {}
    for category, labs in lab_categories.items():
        lab_utilization[category] = {
            "total_slots": len(labs) * len(time_slots) * 5,  # 5 days
            "used_slots": 0,
            "labs": {}
        }
        
        for lab_name in labs:
            # Find matching subjects/rooms
            lab_entries = []
            for entry in entries:
                if entry.subject and any(lab.lower() in entry.subject.name.lower() for lab in labs):
                    lab_entries.append(entry)
                elif entry.room and any(lab.lower() in entry.room.name.lower() for lab in labs):
                    lab_entries.append(entry)
            
            lab_utilization[category]["used_slots"] += len(lab_entries)
            lab_utilization[category]["labs"][lab_name] = {
                "entries": len(lab_entries),
                "utilization": round((len(lab_entries) / (len(time_slots) * 5)) * 100) if len(time_slots) > 0 else 0
            }
    
    # Calculate overall lab utilization
    for category in lab_utilization:
        total = lab_utilization[category]["total_slots"]
        used = lab_utilization[category]["used_slots"]
        lab_utilization[category]["utilization_percentage"] = round((used / total) * 100) if total > 0 else 0
    
    # Department-wise analysis
    departments = ["Computer Engineering", "AI & Data Science", "Electronics", "Mechanical"]
    dept_analysis = {}
    
    for dept in departments:
        dept_entries = [e for e in entries if e.subject and dept.lower() in (e.subject.name.lower() if e.subject else "")]
        dept_analysis[dept] = {
            "total_periods": len(dept_entries),
            "lab_periods": len([e for e in dept_entries if e.subject and (e.subject.is_lab or 'lab' in e.subject.name.lower())]),
            "theory_periods": len([e for e in dept_entries if e.subject and not (e.subject.is_lab or 'lab' in e.subject.name.lower())])
        }
    
    return {
        "lab_utilization": lab_utilization,
        "department_analysis": dept_analysis,
        "time_slots": time_slots,
        "total_lab_entries": len([e for e in entries if e.subject and (e.subject.is_lab or 'lab' in e.subject.name.lower())]),
        "summary": {
            "total_entries": len(entries),
            "lab_percentage": round((len([e for e in entries if e.subject and (e.subject.is_lab or 'lab' in e.subject.name.lower())]) / len(entries)) * 100) if entries else 0
        }
    }

@router.post("/workloads", response_model=schemas.TeacherWorkloadAspect)
def create_custom_workload(
    workload: schemas.TeacherWorkloadAspectCreate, 
    db: Session = Depends(get_db)
):
    """
    Create a custom workload aspect for a teacher.
    """
    db_workload = models.TeacherWorkloadAspect(**workload.model_dump())
    db.add(db_workload)
    db.commit()
    db.refresh(db_workload)
    return db_workload

@router.delete("/workloads/{workload_id}")
def delete_custom_workload(workload_id: int, db: Session = Depends(get_db)):
    """
    Delete a custom workload aspect for a teacher.
    """
    db_workload = db.query(models.TeacherWorkloadAspect).filter(
        models.TeacherWorkloadAspect.id == workload_id
    ).first()
    
    if not db_workload:
        raise HTTPException(status_code=404, detail="Workload Aspect not found")
        
    db.delete(db_workload)
    db.commit()
    return {"status": "success"}

@router.post("/workload-override", response_model=schemas.WorkloadOverride)
def save_workload_override(
    override: schemas.WorkloadOverrideCreate,
    db: Session = Depends(get_db)
):
    """
    Save or update a full JSON override snapshot of the teacher workload report.
    version_id=0 is treated as standalone (stored as NULL).
    """
    actual_version_id = None if override.version_id == 0 else override.version_id
    
    if actual_version_id is None:
        db_override = db.query(models.WorkloadReportOverride).filter(
            models.WorkloadReportOverride.version_id == None
        ).first()
    else:
        db_override = db.query(models.WorkloadReportOverride).filter(
            models.WorkloadReportOverride.version_id == actual_version_id
        ).first()
    
    if db_override:
        db_override.report_data = override.report_data
    else:
        db_override = models.WorkloadReportOverride(
            version_id=actual_version_id,
            report_data=override.report_data
        )
        db.add(db_override)
        
    db.commit()
    db.refresh(db_override)
    return db_override

@router.delete("/workload-override/{version_id}")
def delete_workload_override(version_id: int, db: Session = Depends(get_db)):
    """
    Delete the JSON override snapshot, reverting to dynamically calculated data.
    version_id=0 is treated as standalone (stored as NULL).
    """
    actual_version_id = None if version_id == 0 else version_id
    
    if actual_version_id is None:
        db_override = db.query(models.WorkloadReportOverride).filter(
            models.WorkloadReportOverride.version_id == None
        ).first()
    else:
        db_override = db.query(models.WorkloadReportOverride).filter(
            models.WorkloadReportOverride.version_id == actual_version_id
        ).first()
    
    if not db_override:
        return {"status": "success", "message": "No override existed."}
        
    db.delete(db_override)
    db.commit()
    return {"status": "success", "message": "Override cleared."}

@router.post("/workload-from-docx/{version_id}")
async def upload_workload_docx(
    version_id: int, 
    file: UploadFile = File(...), 
    db: Session = Depends(get_db)
):
    """
    Parse a Word document (.docx) containing a Load Factor / Load Distribution table
    and save it as a JSON override.
    
    Handles complex tables with:
    - Merged cells (same Sr No = same teacher across rows)
    - Multi-line faculty names with designation (e.g. "Dr. Name-\\nAssociate Professor")
    - Complex practical formats (SL1, CL2, MP, "2 * 6 (A-4,B-2)")
    - Empty separator rows, header tables, page titles
    - FE entries with blank faculty
    """
    content = await file.read()
    
    # Debug dump
    try:
        with open("last_uploaded_debug.docx", "wb") as f:
            f.write(content)
    except:
        pass
        
    try:
        doc = docx.Document(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Invalid .docx file: {str(e)}")
    
    if not doc.tables:
        raise HTTPException(status_code=400, detail="No tables found in the document")
    
    # ── Find the data table (the one with 8 columns: Sr, Faculty, Course, Class, Theory, Practical, Total, Project)
    data_table = None
    for tbl in doc.tables:
        if len(tbl.columns) >= 7:
            data_table = tbl
            break
    
    if data_table is None:
        # Fallback: use the last table (often the biggest)
        data_table = doc.tables[-1]
    
    print(f"DOCX PARSER: Using table with {len(data_table.rows)} rows, {len(data_table.columns)} cols")
    
    # ── Parse all rows
    teacher_load = []         # Final output list
    teacher_by_sr = {}        # Track teachers by Sr No to merge rows
    current_teacher = None    # Active teacher for rows without Sr No
    skipped_rows = 0
    courses_parsed = 0
    
    SKIP_KEYWORDS = [
        "d.y. patil", "load distribution", "page", "academic year",
        "autonomous", "semester"
    ]
    HEADER_KEYWORDS = ["sr no", "sr. no", "faculty", "theory", "practical", "total load", "course load"]
    
    for ri, row in enumerate(data_table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        
        # Ensure we have enough columns
        while len(cells) < 8:
            cells.append("")
        
        cell_text_lower = " ".join(cells).lower()
        
        # ── Skip title / page-break rows
        if any(kw in cell_text_lower for kw in SKIP_KEYWORDS):
            skipped_rows += 1
            continue
            
        # ── Skip header rows
        if any(kw in cell_text_lower for kw in HEADER_KEYWORDS):
            skipped_rows += 1
            continue
        
        # ── Skip completely empty rows
        if all(c == "" for c in cells):
            skipped_rows += 1
            continue
        
        sr_no = cells[0].strip()
        faculty_raw = cells[1].strip()
        course_name = cells[2].strip()
        class_div = cells[3].strip()
        theory = cells[4].strip() or "-"
        practical = cells[5].strip() or "-"
        total_load_str = cells[6].strip()
        project = cells[7].strip() or "-"
        
        # Skip rows with no course name (probably artifacts)
        if not course_name:
            skipped_rows += 1
            continue
        
        # ── Parse faculty name and designation
        faculty_name = ""
        designation = ""
        
        if faculty_raw:
            # Split on newline to separate name from designation
            lines = [l.strip() for l in faculty_raw.split('\n') if l.strip()]
            
            if lines:
                raw_name = lines[0]
                
                # Handle trailing dash: "Dr. Name-" or "Dr. Name- HOD"
                if raw_name.endswith('-'):
                    raw_name = raw_name[:-1].strip()
                    if len(lines) > 1:
                        designation = lines[1]
                elif '- ' in raw_name:
                    # "Dr. V. G. Kottawar- HOD"
                    parts = raw_name.rsplit('- ', 1)
                    raw_name = parts[0].strip()
                    designation = parts[1].strip()
                elif '-' in raw_name and raw_name.count('-') == 1:
                    # "Mrs. Snehal Malpani- Assistant Professor"  
                    idx = raw_name.rfind('-')
                    possible_desig = raw_name[idx+1:].strip()
                    # Only split if the part after dash looks like a designation
                    desig_keywords = ['professor', 'hod', 'assistant', 'associate', 'lecturer', 'instructor']
                    if any(kw in possible_desig.lower() for kw in desig_keywords):
                        raw_name = raw_name[:idx].strip()
                        designation = possible_desig
                    elif len(lines) > 1:
                        designation = lines[1]
                elif len(lines) > 1:
                    designation = lines[1]
                
                faculty_name = raw_name.strip()
        
        # ── Determine if this is a new teacher or continuation of previous
        is_new_teacher = False
        
        if sr_no and sr_no.isdigit():
            # Check if we already have a teacher with this Sr No
            if sr_no in teacher_by_sr:
                current_teacher = teacher_by_sr[sr_no]
            else:
                is_new_teacher = True
        elif faculty_raw and current_teacher:
            # Faculty name present but no new Sr No — check if it matches current
            if faculty_raw != current_teacher.get("_raw_faculty", ""):
                is_new_teacher = True
        elif not faculty_raw and current_teacher:
            # No faculty name, continue with current teacher
            pass
        elif faculty_raw:
            # Has faculty name but no Sr No — treat as new
            is_new_teacher = True
        else:
            # No faculty, no Sr No — create an anonymous entry
            if not current_teacher:
                is_new_teacher = True
                faculty_name = course_name  # Use course name as placeholder
        
        if is_new_teacher:
            # Parse total load for the teacher
            total_periods = 0
            if total_load_str:
                match = re.search(r'\d+', total_load_str)
                if match:
                    total_periods = int(match.group())
            
            current_teacher = {
                "_raw_faculty": faculty_raw,
                "name": faculty_name or course_name or f"Unknown Teacher",
                "email": "",
                "designation": designation,
                "total_periods": total_periods,
                "courses_detail": []
            }
            teacher_load.append(current_teacher)
            
            if sr_no and sr_no.isdigit():
                teacher_by_sr[sr_no] = current_teacher
        else:
            # Update total_periods if this row has a higher value
            if total_load_str and current_teacher:
                match = re.search(r'\d+', total_load_str)
                if match:
                    new_total = int(match.group())
                    if new_total > current_teacher.get("total_periods", 0):
                        current_teacher["total_periods"] = new_total
        
        # ── Add course detail to current teacher
        if current_teacher is not None:
            current_teacher["courses_detail"].append({
                "course_name": course_name,
                "class_division": class_div,
                "theory_hours": theory,
                "practical_hours": practical,
                "project_hours": project,
                "total_load": 0,
                "is_custom": True
            })
            courses_parsed += 1
    
    # ── Clean up internal fields
    for t in teacher_load:
        t.pop("_raw_faculty", None)
        # Generate email from name if missing
        if not t.get("email"):
            name_parts = t["name"].lower()
            name_parts = name_parts.replace('dr.', '').replace('mrs.', '').replace('mr.', '').replace('ms.', '')
            name_parts = name_parts.replace('prof.', '').strip().split()
            t["email"] = '.'.join(name_parts) + '@college.edu' if name_parts else 'unknown@college.edu'
    
    if not teacher_load:
        raise HTTPException(
            status_code=400, 
            detail=f"Could not extract any valid teacher rows from the document. "
                   f"Found {len(data_table.rows)} rows but all were skipped. "
                   f"Table has {len(data_table.columns)} columns (expected 8)."
        )

    # ── Save to DB
    actual_version_id = None if version_id == 0 else version_id
    
    if actual_version_id is None:
        db_override = db.query(models.WorkloadReportOverride).filter(
            models.WorkloadReportOverride.version_id == None
        ).first()
    else:
        db_override = db.query(models.WorkloadReportOverride).filter(
            models.WorkloadReportOverride.version_id == actual_version_id
        ).first()
    
    if db_override:
        db_override.report_data = teacher_load
    else:
        db_override = models.WorkloadReportOverride(
            version_id=actual_version_id,
            report_data=teacher_load
        )
        db.add(db_override)
        
    db.commit()
    return {
        "status": "success", 
        "message": f"Parsed and saved from DOCX: {len(teacher_load)} teachers, {courses_parsed} courses ({skipped_rows} rows skipped).",
        "teachers_parsed": len(teacher_load),
        "courses_parsed": courses_parsed,
        "skipped_rows": skipped_rows,
        "data": teacher_load
    }

