"""
Direct DB import from DOCX - uses SQLAlchemy to bypass API limitations.
Clears all existing data and repopulates from the DOCX file.
"""
import sys
import os
import re
import docx
import io

# Add project paths
sys.path.insert(0, os.path.dirname(__file__))
from app.infrastructure.database import SessionLocal, engine
from app.infrastructure import models

def parse_docx(filepath):
    """Parse DOCX and return structured data"""
    doc = docx.Document(filepath)
    
    # Find data table (7+ columns)
    data_table = None
    for tbl in doc.tables:
        if len(tbl.columns) >= 7:
            data_table = tbl
            break
    if not data_table:
        data_table = doc.tables[-1]
    
    SKIP_KW = ["d.y. patil", "load distribution", "page", "academic year", "autonomous", "semester"]
    HEADER_KW = ["sr no", "sr. no", "faculty", "theory", "practical", "total load", "course load"]
    
    teachers = []
    teacher_by_sr = {}
    current_teacher = None
    
    for row in data_table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        while len(cells) < 8:
            cells.append("")
        
        combined = " ".join(cells).lower()
        if any(kw in combined for kw in SKIP_KW) or any(kw in combined for kw in HEADER_KW):
            continue
        if all(c == "" for c in cells):
            continue
        
        sr_no, faculty_raw, course, cls_div = cells[0], cells[1], cells[2].replace("\n", " ").strip(), cells[3]
        theory, practical, total_str, project = cells[4] or "-", cells[5] or "-", cells[6], cells[7] or "-"
        
        if not course:
            continue
        
        # Parse faculty name
        faculty_name, designation = "", ""
        if faculty_raw:
            lines = [l.strip() for l in faculty_raw.split('\n') if l.strip()]
            if lines:
                raw = lines[0]
                if raw.endswith('-'):
                    faculty_name = raw[:-1].strip()
                    designation = lines[1] if len(lines) > 1 else ""
                elif '- ' in raw:
                    parts = raw.rsplit('- ', 1)
                    faculty_name = parts[0].strip()
                    designation = parts[1].strip()
                else:
                    faculty_name = raw.strip()
                    if '-' in raw and raw.count('-') == 1:
                        idx = raw.rfind('-')
                        after = raw[idx+1:].strip()
                        if any(kw in after.lower() for kw in ['professor', 'hod', 'assistant', 'associate']):
                            faculty_name = raw[:idx].strip()
                            designation = after
                    if len(lines) > 1 and not designation:
                        designation = lines[1]
        
        # Determine teacher
        is_new = False
        if sr_no and sr_no.isdigit():
            if sr_no in teacher_by_sr:
                current_teacher = teacher_by_sr[sr_no]
            else:
                is_new = True
        elif faculty_raw and (not current_teacher or faculty_raw != current_teacher.get("_raw")):
            is_new = True
        elif not faculty_raw and not current_teacher:
            is_new = True
            faculty_name = course
        
        if is_new:
            total = 0
            if total_str:
                m = re.search(r'\d+', total_str)
                if m: total = int(m.group())
            
            current_teacher = {
                "_raw": faculty_raw,
                "name": faculty_name or course or "Unknown",
                "designation": designation,
                "total_periods": total,
                "courses": []
            }
            teachers.append(current_teacher)
            if sr_no and sr_no.isdigit():
                teacher_by_sr[sr_no] = current_teacher
        elif total_str and current_teacher:
            m = re.search(r'\d+', total_str)
            if m:
                v = int(m.group())
                if v > current_teacher.get("total_periods", 0):
                    current_teacher["total_periods"] = v
        
        if current_teacher:
            # Parse class divisions
            raw_div = cls_div.replace(" ", "")
            class_names = []
            if raw_div and "-" in raw_div:
                parts = raw_div.split("-", 1)
                prefix = parts[0] + "-"
                for s in parts[1].split(","):
                    s = s.strip()
                    if s:
                        class_names.append(prefix + s)
            elif raw_div:
                class_names.append(raw_div)
            
            # Detect lab type
            prac_upper = str(practical).upper()
            is_lab = any(kw in prac_upper for kw in ["SL1","SL2","CL1","CL2","MP"])
            
            current_teacher["courses"].append({
                "name": course,
                "class_names": class_names,
                "theory": theory,
                "practical": practical,
                "project": project,
                "is_lab": is_lab
            })
    
    # Clean up
    for t in teachers:
        t.pop("_raw", None)
    
    return teachers


def import_to_db(docx_path):
    """Import DOCX data into DB"""
    db = SessionLocal()
    
    try:
        # ── Parse DOCX
        print("Parsing DOCX...")
        teachers_data = parse_docx(docx_path)
        # Filter out department entries
        teachers_data = [t for t in teachers_data if t["name"] not in ["T & P", "HASS"]]
        print(f"  Found {len(teachers_data)} teachers\n")
        
        # ── Collect all unique subjects and classes
        all_subjects = {}  # name -> is_lab
        all_classes = set()
        
        for t in teachers_data:
            for c in t["courses"]:
                name = c["name"]
                if name not in all_subjects:
                    all_subjects[name] = c["is_lab"]
                elif c["is_lab"]:
                    all_subjects[name] = True
                for cls in c["class_names"]:
                    all_classes.add(cls)
        
        print(f"  Unique subjects: {len(all_subjects)}")
        print(f"  Unique classes: {len(all_classes)}\n")
        
        # ── Clear existing data (order matters for FK constraints)
        print("Clearing existing data...")
        db.query(models.TimetableEntry).delete()
        db.query(models.TimetableVersion).delete()
        # Clear lesson associations
        db.execute(models.lesson_teachers.delete())
        db.execute(models.lesson_class_groups.delete())
        db.execute(models.lesson_subjects.delete())
        db.query(models.Lesson).delete()
        db.query(models.Subject).delete()
        db.query(models.TeacherWorkloadAspect).delete()
        db.query(models.Teacher).delete()
        db.query(models.ClassGroup).delete()
        db.query(models.Room).delete()
        db.query(models.Substitution).delete()
        db.commit()
        print("  Done\n")
        
        # ── Insert Teachers
        print("Inserting teachers...")
        teacher_db_map = {}  # name -> Teacher model
        for t in teachers_data:
            name = t["name"]
            # Generate email
            cleaned = name.lower()
            for prefix in ['dr.', 'mrs.', 'mr.', 'ms.', 'prof.']:
                cleaned = cleaned.replace(prefix, '')
            email_parts = cleaned.strip().split()
            email = '.'.join(email_parts) + '@college.edu'
            
            max_hours = t.get("total_periods", 20)
            if max_hours < 5:
                max_hours = 20
            
            teacher = models.Teacher(
                name=name,
                email=email,
                max_hours_per_week=max_hours,
                available_slots=[]
            )
            db.add(teacher)
            db.flush()
            teacher_db_map[name] = teacher
            print(f"  + {name} (id={teacher.id}, max={max_hours}h)")
        
        # ── Insert Class Groups
        print(f"\nInserting {len(all_classes)} class groups...")
        class_db_map = {}
        for cls_name in sorted(all_classes):
            cg = models.ClassGroup(name=cls_name, student_count=60)
            db.add(cg)
            db.flush()
            class_db_map[cls_name] = cg
            print(f"  + {cls_name} (id={cg.id})")
        
        # ── Insert Rooms (one per class + 3 labs)
        print(f"\nInserting rooms...")
        room_num = 1
        for cls_name in sorted(all_classes):
            room = models.Room(name=f"Room-{room_num:02d}", capacity=65, type="LectureHall", resources=[])
            db.add(room)
            print(f"  + Room-{room_num:02d} (LectureHall)")
            room_num += 1
        
        for i in range(1, 4):
            lab = models.Room(name=f"Lab-{i:02d}", capacity=30, type="Lab", resources=["Computers", "Projector"])
            db.add(lab)
            print(f"  + Lab-{i:02d}")
        
        # ── Insert Subjects
        print(f"\nInserting {len(all_subjects)} subjects...")
        subject_db_map = {}
        seen_codes = set()
        
        # Find first teacher for each subject
        subject_first_teacher = {}
        for t in teachers_data:
            for c in t["courses"]:
                if c["name"] not in subject_first_teacher:
                    subject_first_teacher[c["name"]] = t["name"]
        
        for subj_name, is_lab in sorted(all_subjects.items()):
            # Generate unique code
            words = subj_name.split()
            code = ''.join(w[0].upper() for w in words if len(w) > 0 and w[0].isalpha())[:6]
            if not code:
                code = subj_name[:4].upper()
            orig = code
            counter = 1
            while code in seen_codes:
                code = orig + str(counter)
                counter += 1
            seen_codes.add(code)
            
            # Get teacher_id
            teacher_name = subject_first_teacher.get(subj_name)
            teacher_id = teacher_db_map[teacher_name].id if teacher_name and teacher_name in teacher_db_map else None
            
            subj = models.Subject(
                name=subj_name,
                code=code,
                is_lab=is_lab,
                credits=2 if is_lab else 3,
                required_room_type="Lab" if is_lab else "LectureHall",
                duration_slots=2 if is_lab else 1,
                teacher_id=teacher_id
            )
            db.add(subj)
            db.flush()
            subject_db_map[subj_name] = subj
            lab_str = " [LAB]" if is_lab else ""
            print(f"  + {subj_name} ({code}{lab_str})")
        
        # ── Insert Lessons
        print(f"\nCreating lessons...")
        lessons_created = 0
        lesson_keys_seen = set()
        
        for t in teachers_data:
            teacher = teacher_db_map.get(t["name"])
            if not teacher:
                continue
            
            for c in t["courses"]:
                subj = subject_db_map.get(c["name"])
                if not subj:
                    continue
                
                for cls_name in c["class_names"]:
                    cg = class_db_map.get(cls_name)
                    if not cg:
                        continue
                    
                    key = (teacher.id, subj.id, cg.id)
                    if key in lesson_keys_seen:
                        continue
                    lesson_keys_seen.add(key)
                    
                    lesson = models.Lesson(
                        lessons_per_week=1,
                        length_per_lesson=2 if c["is_lab"] else 1
                    )
                    db.add(lesson)
                    db.flush()
                    
                    # Add associations
                    db.execute(models.lesson_teachers.insert().values(
                        lesson_id=lesson.id, teacher_id=teacher.id
                    ))
                    db.execute(models.lesson_class_groups.insert().values(
                        lesson_id=lesson.id, class_group_id=cg.id
                    ))
                    db.execute(models.lesson_subjects.insert().values(
                        lesson_id=lesson.id, subject_id=subj.id
                    ))
                    lessons_created += 1
        
        print(f"  Created {lessons_created} lessons")
        
        # ── Ensure time slots exist
        print(f"\nChecking time slots...")
        slot_count = db.query(models.TimeSlot).count()
        if slot_count < 25:
            print("  Creating time slots...")
            # Clear existing
            db.query(models.TimeSlot).delete()
            
            days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
            periods = [
                (1, "09:00", "09:50", False),
                (2, "09:50", "10:40", False),
                (3, "10:40", "11:30", False),
                (4, "11:40", "12:30", False),  # After short break
                (5, "12:30", "13:20", False),
                (6, "14:00", "14:50", False),  # After lunch
                (7, "14:50", "15:40", False),
            ]
            
            for day in days:
                for period, start, end, is_break in periods:
                    slot = models.TimeSlot(
                        day=day, period=period,
                        start_time=start, end_time=end,
                        is_break=is_break
                    )
                    db.add(slot)
            
            db.flush()
            new_count = db.query(models.TimeSlot).count()
            print(f"  Created {new_count} time slots (5 days × 7 periods)")
        else:
            print(f"  {slot_count} time slots already exist")
        
        # ── Commit everything
        db.commit()
        
        # ── Also update the workload override with the parsed data
        print(f"\nUpdating workload override...")
        teacher_load_override = []
        for t in teachers_data:
            teacher_load_override.append({
                "name": t["name"],
                "designation": t.get("designation", ""),
                "email": teacher_db_map[t["name"]].email if t["name"] in teacher_db_map else "",
                "total_periods": t.get("total_periods", 0),
                "courses_detail": [
                    {
                        "course_name": c["name"],
                        "class_division": ", ".join(c["class_names"]),
                        "theory_hours": c["theory"],
                        "practical_hours": c["practical"],
                        "project_hours": c["project"],
                        "total_load": 0,
                        "is_custom": True
                    } for c in t["courses"]
                ]
            })
        
        # Save override with NULL version_id (standalone)
        override = db.query(models.WorkloadReportOverride).filter(
            models.WorkloadReportOverride.version_id == None
        ).first()
        if override:
            override.report_data = teacher_load_override
        else:
            override = models.WorkloadReportOverride(
                version_id=None,
                report_data=teacher_load_override
            )
            db.add(override)
        db.commit()
        print("  Done")
        
        # ── Final summary
        print("\n" + "=" * 60)
        print("IMPORT COMPLETE!")
        print("=" * 60)
        print(f"  Teachers:     {len(teacher_db_map)}")
        print(f"  Subjects:     {len(subject_db_map)}")
        print(f"  Classes:      {len(class_db_map)}")
        print(f"  Rooms:        {len(all_classes) + 3}")
        print(f"  Lessons:      {lessons_created}")
        print(f"  Time slots:   {db.query(models.TimeSlot).count()}")
        print(f"\nReady to generate timetable!")
        
    except Exception as e:
        db.rollback()
        print(f"\nERROR: {e}")
        import traceback
        traceback.print_exc()
    finally:
        db.close()


if __name__ == "__main__":
    import_to_db("last_uploaded_debug.docx")
